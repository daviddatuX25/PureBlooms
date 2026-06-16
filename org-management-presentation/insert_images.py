from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document('/home/user/Downloads/Pure Blooms/org-management-presentation/PureBlooms_Developed_System_Documentation.docx')

# Image mappings: paragraph index -> (image_path, caption)
images = {
    73: ("/home/user/Downloads/Pure Blooms/org-management-presentation/diagrams/mvc_architecture.png",
         "Figure 1. MVC Architecture Diagram"),
    93: ("/home/user/Downloads/Pure Blooms/org-management-presentation/screenshots/customer_storefront_homepage.png",
         "Figure 2. Customer homepage displaying featured bouquets, category tabs, and header navigation."),
    99: ("/home/user/Downloads/Pure Blooms/org-management-presentation/screenshots/customer_storefront_homepage.png",
         "Figure 3. Checkout screen displaying customer details form, payment method toggles, and dynamic add-on selection. [CHECKOUT SCREENSHOT NEEDED]"),
    104: ("/home/user/Downloads/Pure Blooms/org-management-presentation/screenshots/customer_order_tracking.png",
         "Figure 4. Customer order tracking page displaying a visual vertical timeline of the order's fulfillment path."),
    112: ("/home/user/Downloads/Pure Blooms/org-management-presentation/screenshots/admin_dashboard.png",
         "Figure 5. Admin Dashboard showing high-level KPIs, metric cards, and low-stock notification alerts."),
    117: ("/home/user/Downloads/Pure Blooms/org-management-presentation/screenshots/admin_order_pipeline.png",
         "Figure 6. Admin order list interface showing filter tabs, action buttons, and status badge indicators."),
    124: ("/home/user/Downloads/Pure Blooms/org-management-presentation/screenshots/admin_sales_report.png",
         "Figure 7. Admin Sales Report displaying order frequencies and total revenue metrics."),
    129: ("/home/user/Downloads/Pure Blooms/org-management-presentation/screenshots/admin_settings_maintenance.png",
         "Figure 8. System Settings page displaying payment toggles and the global Maintenance Mode switch."),
    157: ("/home/user/Downloads/Pure Blooms/org-management-presentation/diagrams/om_framework.png",
         "Figure 9. OM Framework Diagram - Mapping Organizational Management Functions to PureBlooms Modules"),
    178: ("/home/user/Downloads/Pure Blooms/org-management-presentation/diagrams/org_chart.png",
         "Figure 10. Team Organizational Structure - Agile IT Consultancy Firm"),
}

# Process in reverse order so indices don't shift
sorted_indices = sorted(images.keys(), reverse=True)

for idx in sorted_indices:
    img_path, caption_text = images[idx]

    if not os.path.exists(img_path):
        print(f"WARNING: Image not found: {img_path}")
        continue

    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]

        # Clear the paragraph
        for run in p.runs:
            run.text = ""

        # Add the image
        run = p.add_run()
        run.add_picture(img_path, width=Inches(6.5))

        # Center the image
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add caption as a new paragraph AFTER using XML
        from docx.oxml import OxmlElement
        new_p = OxmlElement('w:p')
        p._element.addnext(new_p)
        
        # Now we need to get the new paragraph object
        # The new paragraph is now in doc.paragraphs, but at the end
        # We'll just create a new paragraph object and format it
        from docx.text.paragraph import Paragraph
        caption_p = Paragraph(new_p, doc)
        
        caption_run = caption_p.add_run(caption_text)
        caption_run.font.size = Pt(10)
        caption_run.font.italic = True
        caption_run.font.color.rgb = RGBColor(89, 89, 89)
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        print(f"Inserted at paragraph {idx}: {os.path.basename(img_path)}")
    else:
        print(f"ERROR: Index {idx} out of range")

# Save
output_path = "/home/user/Downloads/Pure Blooms/org-management-presentation/PureBlooms_Developed_System_Documentation_with_images.docx"
doc.save(output_path)
print(f"\nSaved to: {output_path}")
